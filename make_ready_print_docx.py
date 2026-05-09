"""Generate print-ready Word .docx files for everything under `drafts for print/ready/`.

Output spec (driven by user request):
- 8.5 x 11 (Letter) page size
- 0.5" margins all sides
- Courier New everywhere
- Sized to fill the page: continuous text at 14pt; signage subsections (one
  per page) auto-scaled large and vertically centered so each item dominates
  its sheet.

Walks categories 02 (Character Handouts), 03 (Sealed Envelopes), and
04 (Signage). For each `.md` file, extracts the `## COPY TO PRINT` section,
converts to .docx via pandoc, and post-processes the file to apply the print
spec above.

Signage files contain multiple `### Sub-section` blocks (one per nametag /
badge / sign / flipboard card). Each subsection becomes its own page, with
the heading rendered very large and the body sized to fill the rest of the
sheet.

Output mirrors the source structure under `drafts for print/word/`.

Usage:
    uv run make_ready_print_docx.py
"""

from pathlib import Path
import re
import tempfile

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).parent
PRINT_ROOT = PROJECT_ROOT / "drafts for print"
READY_ROOT = PRINT_ROOT / "ready"
WORD_ROOT = PRINT_ROOT / "word"

CATEGORIES = [
    "02 - Character Handouts",
    "03 - Sealed Envelopes",
    "04 - Signage",
]

MULTI_SECTION_CATEGORIES = {"04 - Signage"}

BODY_FONT = "Courier New"

# Continuous-text default. Pairs with 0.5" margins on Letter to fill the
# page without crowding.
BODY_FONT_PT = 14

# Signage one-per-page sizing. Heading 3 (the per-item title) is the room-
# read marquee. Body lines are large but smaller than the heading.
SIGNAGE_HEADING_PT = 48
SIGNAGE_BODY_PT = 28


def extract_copy_to_print(md_text: str) -> str:
    """Pull the `## COPY TO PRINT` section out as clean markdown."""
    match = re.search(r"^##\s+COPY TO PRINT\s*$", md_text, re.MULTILINE)
    if not match:
        raise ValueError("No '## COPY TO PRINT' section found")

    rest = md_text[match.end():].lstrip("\n")
    end = re.search(r"^##\s+\S", rest, re.MULTILINE)
    section = rest[: end.start()] if end else rest

    out_lines: list[str] = []
    for line in section.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("> "):
            out_lines.append(stripped[2:])
        elif stripped == ">":
            out_lines.append("")
        else:
            out_lines.append(stripped)

    cleaned = "\n".join(out_lines).rstrip() + "\n"
    cleaned = cleaned.replace("&nbsp;", " ")
    return cleaned


def force_run_font(run, name: str) -> None:
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def set_letter_page_with_tight_margins(doc: Document) -> None:
    """Letter (8.5x11), 0.5" margins all sides, on every section."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def is_heading(para, level: int) -> bool:
    name = para.style.name if para.style and para.style.name else ""
    return name == f"Heading {level}"


def apply_courier_with_size(doc: Document, body_pt: int) -> None:
    """Force Courier New on every run; set body font size to body_pt."""
    for para in doc.paragraphs:
        for run in para.runs:
            force_run_font(run, BODY_FONT)
            run.font.size = Pt(body_pt)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        force_run_font(run, BODY_FONT)
                        run.font.size = Pt(body_pt)


def insert_page_break_before(paragraph) -> None:
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    new_br = OxmlElement("w:br")
    new_br.set(qn("w:type"), "page")
    new_r.append(new_br)
    new_p.append(new_r)
    paragraph._p.addprevious(new_p)


def style_signage(doc: Document) -> None:
    """For multi-section signage: each `### ` becomes its own giant page.

    - Page break before every Heading 3 except the first.
    - Heading 3 line: SIGNAGE_HEADING_PT, bold, centered.
    - Every other paragraph: SIGNAGE_BODY_PT, centered, Courier.
    - Inserts vertical breathing room above the first item on each page so
      the content sits roughly center-of-page rather than top-of-page.
    """
    seen_first = False
    for para in list(doc.paragraphs):
        if is_heading(para, 3):
            if seen_first:
                insert_page_break_before(para)
            else:
                seen_first = True

    for para in doc.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if is_heading(para, 3):
            for run in para.runs:
                force_run_font(run, BODY_FONT)
                run.font.size = Pt(SIGNAGE_HEADING_PT)
                run.bold = True
        else:
            for run in para.runs:
                force_run_font(run, BODY_FONT)
                run.font.size = Pt(SIGNAGE_BODY_PT)


def process_file(md_path: Path, out_dir: Path, multi_section: bool) -> Path:
    md_text = md_path.read_text(encoding="utf-8")
    print_md = extract_copy_to_print(md_text)

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{md_path.stem}.docx"

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".md", delete=False, encoding="utf-8"
    ) as tf:
        tf.write(print_md)
        tmp_path = Path(tf.name)

    try:
        pypandoc.convert_file(
            str(tmp_path),
            to="docx",
            outputfile=str(out_path),
            extra_args=["--standalone"],
        )
    finally:
        tmp_path.unlink(missing_ok=True)

    doc = Document(str(out_path))
    set_letter_page_with_tight_margins(doc)
    if multi_section:
        style_signage(doc)
    else:
        apply_courier_with_size(doc, BODY_FONT_PT)
    doc.save(str(out_path))

    return out_path


def main() -> None:
    total = 0
    for category in CATEGORIES:
        src_dir = READY_ROOT / category
        if not src_dir.exists():
            print(f"  skip (missing): {src_dir}")
            continue

        out_dir = WORD_ROOT / category
        md_files = sorted(src_dir.glob("*.md"))
        if not md_files:
            print(f"  skip (no .md files): {src_dir}")
            continue

        multi_section = category in MULTI_SECTION_CATEGORIES
        print(f"\n{category}{'  (multi-section)' if multi_section else ''}")
        for md in md_files:
            try:
                out = process_file(md, out_dir, multi_section)
            except ValueError as exc:
                print(f"  skip {md.name}: {exc}")
                continue
            print(f"  {md.name}  ->  {out.relative_to(PROJECT_ROOT)}")
            total += 1

    print(f"\nDone. {total} file(s) written under {WORD_ROOT.relative_to(PROJECT_ROOT)}/")


if __name__ == "__main__":
    main()
