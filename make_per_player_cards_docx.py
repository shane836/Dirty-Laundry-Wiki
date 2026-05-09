"""Generate 2-cards-per-sheet round-card .docx files, one per character.

Each player gets TWO physical letter-size sheets (4 pages total):
  • SHEET 1 — Round 1 card (LEFT half) + Round 2 card (RIGHT half)
      - Page 1 BACK:  R1 content (left) | R2 content (right)
      - Page 2 FRONT: R2 stop notice (left) | R1 stop notice (right) — mirrored
  • SHEET 2 — Round 3 card (LEFT half) + Round 4 card (RIGHT half)
      - Page 3 BACK:  R3 content (left) | R4 content (right)
      - Page 4 FRONT: R4 stop notice (left) | R3 stop notice (right) — mirrored

After printing duplex (long-edge binding — the consumer-printer default),
cut each sheet vertically at the midline (4.25") to get 2 cards per sheet
(4 cards per player total). Each card is 4.25" wide x 11" tall, with the
stop notice on the front and round content on the back.

If the cards come out misaligned after cutting, your printer is using
short-edge flip — re-print with "long-edge binding" selected.

Usage:
    uv run make_per_player_cards_docx.py            # all characters
    uv run make_per_player_cards_docx.py 1          # just character #01
    uv run make_per_player_cards_docx.py 1 2 5      # characters 1, 2, 5
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).parent
SOURCE_DIR = PROJECT_ROOT / "drafts for print" / "06 - Per-Player Round Cards"
OUTPUT_DIR = PROJECT_ROOT / "docx" / "per-player-cards"
FONT_NAME = "Courier New"

# Round metadata: (label, time, title)
ROUNDS = [
    ("ROUND 1", "8:00 PM – 10:00 PM", "The After-Party"),
    ("ROUND 2", "10:00 PM – 11:00 PM", "Everyone's a Suspect"),
    ("ROUND 3", "11:00 PM – Midnight", "The Investigation"),
    ("ROUND 4", "Midnight", "The Reveal"),
]


# ───────────────────── docx helpers ─────────────────────


def force_run_font(run, size_pt: float | None = None) -> None:
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def remove_default_paragraph(cell) -> None:
    for p in list(cell.paragraphs):
        if not p.text.strip() and not p.runs:
            p._element.getparent().remove(p._element)
            return


def add_paragraph(
    cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: float = 9,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    after_pt: float = 1,
    before_pt: float = 0,
    indent_in: float = 0.0,
) -> None:
    para = cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(after_pt)
    para.paragraph_format.space_before = Pt(before_pt)
    if indent_in:
        para.paragraph_format.left_indent = Inches(indent_in)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        force_run_font(run, size)


def add_runs(
    cell,
    segments,
    *,
    size: float = 9,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    after_pt: float = 1,
    before_pt: float = 0,
    indent_in: float = 0.0,
) -> None:
    """Add a paragraph with multiple runs (for inline bold/italic)."""
    para = cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(after_pt)
    para.paragraph_format.space_before = Pt(before_pt)
    if indent_in:
        para.paragraph_format.left_indent = Inches(indent_in)
    for text, bold, italic in segments:
        if not text:
            continue
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        force_run_font(run, size)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def parse_inline(text: str):
    segments = []
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            segments.append((chunk[2:-2], True, False))
        elif chunk.startswith("*") and chunk.endswith("*"):
            segments.append((chunk[1:-1], False, True))
        else:
            segments.append((chunk, False, False))
    return segments


def set_cell_margins(cell, top=0.06, bottom=0.06, left=0.10, right=0.10) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(int(val * 1440)))  # twips
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def set_table_borders(table, color: str = "999999") -> None:
    """Light dashed cut-line borders so the cutter sees where to slice."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "dashed")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def setup_page(section) -> None:
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)


def add_1x2_table(doc: Document, renderers: list[Callable]) -> None:
    """renderers is [LEFT, RIGHT] — each (cell -> None). Single row, full page height."""
    section = doc.sections[0]
    avail_w = section.page_width - section.left_margin - section.right_margin
    avail_h = section.page_height - section.top_margin - section.bottom_margin
    cell_w = avail_w // 2
    cell_h = avail_h

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False

    row = table.rows[0]
    row.height = cell_h
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    # w:val expects twips (1/20 pt). cell_h is in EMU; convert: 1 twip = 635 EMU.
    trHeight.set(qn("w:val"), str(int(cell_h / 635)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    for c in range(2):
        table.cell(0, c).width = cell_w

    cells = [table.cell(0, 0), table.cell(0, 1)]
    for cell, render in zip(cells, renderers):
        set_cell_margins(cell)
        render(cell)

    set_table_borders(table)


# ───────────────────── card renderers ─────────────────────


def render_stop_notice(cell, round_label: str, time_str: str, title: str, character: str) -> None:
    remove_default_paragraph(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    add_paragraph(cell, "DO NOT READ", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after_pt=0)
    add_paragraph(cell, "UNTIL", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after_pt=4)
    add_paragraph(cell, round_label, bold=True, size=32,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after_pt=14, before_pt=8)
    add_paragraph(cell, time_str, size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after_pt=2)
    add_paragraph(cell, f'"{title}"', italic=True, size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after_pt=24)
    add_paragraph(cell, character.upper(), bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER)


def render_content(cell, round_label: str, time_str: str, title: str,
                   character: str, content_lines: list[str]) -> None:
    remove_default_paragraph(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Header
    add_paragraph(cell, f"{round_label} · {character.upper()}",
                  bold=True, size=8.5, after_pt=0)
    add_runs(cell,
             [(f"{time_str} · ", False, False), (f'"{title}"', False, True)],
             size=7, after_pt=4)

    for raw in content_lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if re.match(r"^-{3,}\s*$", line):
            continue

        # ## / ### headers
        m = re.match(r"^#{2,4}\s+(.*)$", line)
        if m:
            add_paragraph(cell, m.group(1).strip(),
                          bold=True, size=7.5, before_pt=3, after_pt=1)
            continue

        # Bullet
        b = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if b:
            indent = len(b.group(1))
            text = b.group(2)
            level = 1 if indent < 2 else 2
            prefix = "• " if level == 1 else "◦ "
            indent_in = 0.0 if level == 1 else 0.15
            segs = [(prefix, False, False)] + parse_inline(text)
            add_runs(cell, segs, size=7, after_pt=1, indent_in=indent_in)
            continue

        # Plain paragraph
        add_runs(cell, parse_inline(line), size=7, after_pt=2)


# ───────────────────── markdown parsing ─────────────────────


def parse_md(md_path: Path) -> tuple[str, dict[int, list[str]]]:
    text = md_path.read_text()

    # Strip frontmatter
    character = md_path.stem
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end >= 0:
            front = text[3:end]
            text = text[end + 4:].lstrip("\n")
            cm = re.search(r'character:\s*"([^"]+)"', front)
            if cm:
                character = cm.group(1)

    # Split on "# Round N" headers
    parts = re.split(r"^# Round (\d+)[^\n]*\n", text, flags=re.MULTILINE)
    rounds: dict[int, list[str]] = {}
    for i in range(1, len(parts), 2):
        n = int(parts[i])
        body = parts[i + 1]
        # Trim at next round/newpage/disclaimer if any survived
        body = re.split(r"\n# Round \d+", body)[0]
        body = re.split(r"\n\\newpage", body)[0]
        body = re.split(r"\n# \*\*DO NOT READ", body)[0]

        # Drop the leading "### TIME" line if present
        body_lines = body.strip().split("\n")
        if body_lines and body_lines[0].startswith("###"):
            body_lines = body_lines[1:]
            # also drop any blank line right after
            while body_lines and not body_lines[0].strip():
                body_lines = body_lines[1:]
        rounds[n] = body_lines

    return character, rounds


# ───────────────────── document assembly ─────────────────────


def make_card_doc(md_path: Path) -> Document:
    character, rounds = parse_md(md_path)

    doc = Document()
    setup_page(doc.sections[0])

    def content_renderer(round_idx: int):
        label, time_str, title = ROUNDS[round_idx]
        lines = rounds.get(round_idx + 1, [])
        return lambda cell: render_content(cell, label, time_str, title, character, lines)

    def stop_renderer(round_idx: int):
        label, time_str, title = ROUNDS[round_idx]
        return lambda cell: render_stop_notice(cell, label, time_str, title, character)

    # Each table fills the full page (cell_h = avail_h, hRule="exact"), so
    # subsequent tables naturally land on a new page — no explicit page breaks
    # needed (those would create blank pages between content pages).

    # ── SHEET 1 — Round 1 card (left) + Round 2 card (right) ─────────
    # Page 1 BACK: content. LEFT=R1 | RIGHT=R2
    add_1x2_table(doc, [content_renderer(0), content_renderer(1)])
    # Page 2 FRONT: stop notices, mirrored for long-edge flip. LEFT=R2 | RIGHT=R1
    add_1x2_table(doc, [stop_renderer(1), stop_renderer(0)])

    # ── SHEET 2 — Round 3 card (left) + Round 4 card (right) ─────────
    # Page 3 BACK: content. LEFT=R3 | RIGHT=R4
    add_1x2_table(doc, [content_renderer(2), content_renderer(3)])
    # Page 4 FRONT: stop notices, mirrored. LEFT=R4 | RIGHT=R3
    add_1x2_table(doc, [stop_renderer(3), stop_renderer(2)])

    return doc


# ───────────────────── CLI ─────────────────────


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    md_files = sorted(SOURCE_DIR.glob("*.md"))

    if len(sys.argv) > 1:
        wanted = {a.zfill(2) for a in sys.argv[1:]}
        md_files = [m for m in md_files if m.name[:2] in wanted]
        if not md_files:
            print(f"No matching files for: {sorted(wanted)}")
            return

    for md in md_files:
        out = OUTPUT_DIR / f"{md.stem}.docx"
        doc = make_card_doc(md)
        doc.save(str(out))
        print(f"  {md.name}  →  {out.relative_to(PROJECT_ROOT)}")

    print(f"\nDone. {len(md_files)} file(s) written to {OUTPUT_DIR.relative_to(PROJECT_ROOT)}/")


if __name__ == "__main__":
    main()
