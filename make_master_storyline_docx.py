"""Convert the master storyline to .docx and set every glyph to Courier New.

Pandoc handles the markdown → docx conversion. python-docx then walks every
style and run to force Courier New (pandoc itself can't override fonts without
a reference.docx).

Usage:
    uv run make_master_storyline_docx.py
"""

from pathlib import Path

import pypandoc
from docx import Document
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).parent
SOURCE = PROJECT_ROOT / "00-Master-Storyline.md"
OUTPUT_DIR = PROJECT_ROOT / "docx"
FONT_NAME = "Courier New"


def force_font(run, name: str) -> None:
    """Set the font on a run for Latin, East Asian, and complex scripts."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def apply_font_everywhere(doc: Document, name: str) -> None:
    # Update every style
    for style in doc.styles:
        if hasattr(style, "font") and style.font is not None:
            try:
                style.font.name = name
            except (AttributeError, TypeError):
                pass

    # Update every run in body paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            force_font(run, name)

    # Update every run inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        force_font(run, name)


def main() -> None:
    if not SOURCE.exists():
        print(f"Source not found: {SOURCE}")
        return

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUTPUT_DIR / f"{SOURCE.stem}.docx"

    pypandoc.convert_file(
        str(SOURCE),
        to="docx",
        outputfile=str(out),
        extra_args=["--standalone"],
    )

    doc = Document(str(out))
    apply_font_everywhere(doc, FONT_NAME)
    doc.save(str(out))

    print(f"  {SOURCE.name}  →  {out.relative_to(PROJECT_ROOT)}")
    print(f"  Font: {FONT_NAME}")


if __name__ == "__main__":
    main()
