"""Generate Word .docx files for print-shop evidence drafts.

Each draft markdown has a `## COPY TO PRINT` blockquote with the rendered text
plus layout notes describing how it should look. We extract the blockquote,
convert to .docx via pandoc, then apply per-file font/style profiles that
follow the layout notes (hotel letterhead vs. lab/coroner, monospaced report
body vs. serif narrative, etc.).

Output mirrors the source category structure under `drafts for print/word/`.

Usage:
    uv run make_print_evidence_docx.py
"""

from pathlib import Path
import re
import tempfile

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).parent
PRINT_ROOT = PROJECT_ROOT / "drafts for print"
WORD_ROOT = PRINT_ROOT / "word"

CATEGORIES = [
    "01 - Hotel Letterhead Evidence",
]

# Per-file style profiles. Drives body font, code-block font, and whether the
# top header is treated as hotel letterhead (centered, large) or as a plainer
# government/lab header (left-aligned, smaller).
STYLE_PROFILES: dict[str, dict] = {
    # Hotel letterhead set — Round 2/3 reveals from the mentalist folder.
    "01-Key-Card-Log": {
        "body_font": "Times New Roman",
        "code_font": "Courier New",
        "header_style": "hotel_letterhead",
    },
    "02-Incident-Report": {
        "body_font": "Times New Roman",
        "code_font": "Courier New",
        "header_style": "hotel_letterhead",
    },
    "03-Override-Card-Swipe-Log": {
        "body_font": "Times New Roman",
        "code_font": "Courier New",
        "header_style": "hotel_letterhead",
    },
    # Maintenance form — distinct visual style per spec.
    "04-Maintenance-Work-Order": {
        "body_font": "Calibri",
        "code_font": "Courier New",
        "header_style": "hotel_letterhead_small",
    },
    # NOT hotel letterhead — government/lab official documents.
    "05-Chemical-Analysis-Report": {
        "body_font": "Times New Roman",
        "code_font": "Courier New",
        "header_style": "official_plain",
    },
    "06-Autopsy-Finding": {
        "body_font": "Times New Roman",
        "code_font": "Courier New",
        "header_style": "official_plain",
    },
}

DEFAULT_PROFILE = {
    "body_font": "Times New Roman",
    "code_font": "Courier New",
    "header_style": "hotel_letterhead",
}


def extract_copy_to_print(md_text: str) -> str:
    """Pull the `## COPY TO PRINT` blockquote out and return clean markdown.

    Strips the leading `> ` from each line. Stops at the first non-blockquote,
    non-blank line after the heading.
    """
    match = re.search(r"^##\s+COPY TO PRINT\s*$", md_text, re.MULTILINE)
    if not match:
        raise ValueError("No '## COPY TO PRINT' section found")

    rest = md_text[match.end():].lstrip("\n")
    out_lines: list[str] = []
    for line in rest.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("> "):
            out_lines.append(stripped[2:])
        elif stripped == ">":
            out_lines.append("")
        elif stripped == "":
            out_lines.append("")
        else:
            break

    cleaned = "\n".join(out_lines).rstrip() + "\n"
    # &nbsp; → real non-breaking space so it renders cleanly in Word.
    cleaned = cleaned.replace("&nbsp;", " ")
    return cleaned


def force_run_font(run, name: str) -> None:
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def is_code_paragraph(para) -> bool:
    name = para.style.name if para.style and para.style.name else ""
    return name.startswith("Source Code") or name in {
        "Verbatim Char",
        "HTML Preformatted",
        "Code",
    }


def apply_fonts(doc: Document, body_font: str, code_font: str) -> None:
    for para in doc.paragraphs:
        font = code_font if is_code_paragraph(para) else body_font
        for run in para.runs:
            run_style_name = ""
            try:
                if run.style and run.style.name:
                    run_style_name = run.style.name
            except AttributeError:
                pass
            if "Verbatim" in run_style_name or "Code" in run_style_name:
                force_run_font(run, code_font)
            else:
                force_run_font(run, font)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        force_run_font(run, body_font)


def style_header(doc: Document, header_style: str) -> None:
    """Treat the first paragraph as the document's letterhead/header line."""
    if not doc.paragraphs:
        return
    first = doc.paragraphs[0]

    if header_style == "hotel_letterhead":
        first.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = Pt(18)
    elif header_style == "hotel_letterhead_small":
        first.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = Pt(14)
    elif header_style == "official_plain":
        first.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = Pt(14)
    else:
        return

    for run in first.runs:
        run.bold = True
        run.font.size = size


def process_file(md_path: Path, out_dir: Path) -> Path:
    md_text = md_path.read_text(encoding="utf-8")
    print_md = extract_copy_to_print(md_text)

    profile = STYLE_PROFILES.get(md_path.stem, DEFAULT_PROFILE)

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{md_path.stem}.docx"

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".md", delete=False, encoding="utf-8"
    ) as tf:
        tf.write(print_md)
        tmp_path = Path(tf.name)

    try:
        pypandoc.convert_file(
            str(tmp_path),
            to="docx",
            outputfile=str(out_path),
            extra_args=["--standalone"],
        )
    finally:
        tmp_path.unlink(missing_ok=True)

    doc = Document(str(out_path))
    apply_fonts(doc, profile["body_font"], profile["code_font"])
    style_header(doc, profile["header_style"])
    doc.save(str(out_path))

    return out_path


def main() -> None:
    total = 0
    for category in CATEGORIES:
        src_dir = PRINT_ROOT / category
        if not src_dir.exists():
            print(f"  skip (missing): {src_dir}")
            continue

        out_dir = WORD_ROOT / category
        md_files = sorted(src_dir.glob("*.md"))
        if not md_files:
            print(f"  skip (no .md files): {src_dir}")
            continue

        print(f"\n{category}")
        for md in md_files:
            try:
                out = process_file(md, out_dir)
            except ValueError as exc:
                print(f"  skip {md.name}: {exc}")
                continue
            print(f"  {md.name}  →  {out.relative_to(PROJECT_ROOT)}")
            total += 1

    print(f"\nDone. {total} file(s) written under {WORD_ROOT.relative_to(PROJECT_ROOT)}/")


if __name__ == "__main__":
    main()
